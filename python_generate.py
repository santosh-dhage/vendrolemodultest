import subprocess
from docx import Document
from docx.shared import Inches
import os
# def generate_erd_image():
#     # Execute the graph_models command to generate the ERD image
#     subprocess.run(["python", "manage.py", "graph_models", "-a", "-g", "-o", "erd.png"])

# def create_word_document():
#     # Create a new Word document
#     doc = Document()

#     # Add a heading to the document
#     doc.add_heading('Entity-Relationship Diagram', level=1)

#     # Add the ERD image to the document
#     doc.add_picture('erd.png')

#     # Save the Word document
#     doc.save('erd_document.docx')

# def main():
#     generate_erd_image()
#     create_word_document()

# if __name__ == "__main__":
#     main()
def split_erd_image(image_path, output_dir, num_pages):
    # Create a new Word document
    doc = Document()

    # Add a heading to the document
    doc.add_heading('Entity-Relationship Diagram', level=1)

    # Determine the width of the image to fit on the page
    image_width = Inches(6)  # Adjust as needed

    # Calculate the number of images per page
    images_per_page = num_pages

    # Split the ERD image into multiple parts
    for i in range(num_pages):
        # Insert the ERD image into the document
        doc.add_picture(image_path, width=image_width)

        # Add a page break between images
        if i < num_pages - 1:
            doc.add_page_break()

    # Save the Word document
    output_filename = os.path.join(output_dir, 'erd_document.docx')
    doc.save(output_filename)

def main():
    # Path to the ERD image
    image_path = 'erd.png'  # Update with the path to your ERD image

    # Output directory for the Word document
    output_dir = 'D:\ivendsoft_tenant'  # Update with the desired output directory

    # Number of pages to split the image into
    num_pages = 3  # Update with the desired number of pages

    # Split the ERD image and create the Word document
    split_erd_image(image_path, output_dir, num_pages)

if __name__ == "__main__":
    main()